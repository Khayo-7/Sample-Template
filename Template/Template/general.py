import base64
import imgkit
import datetime
from io import BytesIO 
from docx import Document
from xhtml2pdf import pisa
from docx.shared import Inches
from docxtpl import DocxTemplate
from django.template import Context 
from django.contrib import messages 
from django.http import FileResponse
from django.views.generic import FormView
from django.http import HttpResponse, Http404
from django.db import IntegrityError, transaction
from django.db.models import Q, CharField, DateField
from django.contrib.postgres.search import SearchVector
from django.template.loader import render_to_string, get_template 
from django.shortcuts import render, redirect, reverse, get_object_or_404


# Create your views here.
from Template.trans import phrases
from Template.forms import SearchForm
from Template.faker import generate_random_data

## delete
from users.models import User
from myapps.models import *

# General
class General:  

    def __init__(self, arguments):     

        self.app = arguments['app']
        self.title = arguments['title']
        self.success_url = arguments['success_url'] 
        self.instance = arguments['instance']
        self.instance_type = arguments['instance_type']
        self.form = arguments['form']   
        # self.update_form = arguments['update_form']
        self.serializer = arguments['serializer']  
        self.filter = arguments['filter']  

        self.action = arguments['action'] 
        self.page_type = arguments['page_type']
        self.url_suffix = arguments['url_suffix']
        self.view_all_url = arguments['view_all_url']
        # self.request = arguments['request']
        self.form_template = arguments['form_template']
        self.main_template = arguments['main_template']
        self.view_template = arguments['view_template']
        self.pdf_template = arguments['pdf_template']
        self.view_all_template = arguments['view_all_template']
        self.assurance_template = arguments['assurance_template']

    def get_context(self):
        context = {
                'app': self.app, 
                'title': self.title, 
                'instance_type': self.instance_type, 
                'action': self.action, 
                'page_type': self.page_type, 
                'url_suffix': self.url_suffix, 
                'view_all_url': self.view_all_url
            }
        return context

    def index(self, request):  
        context = self.get_context()
        return render(request, self.main_template, context)
 

    def search(self, request):  

        self.action = "Search"  
        self.page_type = "Result Page"  
        self.form = SearchForm 
        instance = self.instance

        form = self.form(request.POST or request.GET or None)
        context = self.get_context()
        context.update({'form': form})   
         
        if request.method == 'POST' or request.method == 'POST':

            if form.is_valid():
                keyword = form.cleaned_data['keyword'] 

                field_names = [field.name for field in instance._meta.fields if isinstance(field, CharField)]# or isinstance(f, DateField)]
                
                queries = Q()
                for query in [Q(**{field_name + "__icontains": keyword}) for field_name in field_names]:
                    queries = queries | query

                queryset = instance.objects.filter(queries).values()
                # queryset = instance.objects.annotate(search=SearchVector(*field_names),).filter(search__icontains=keyword).values()
                
                filter = self.filter(request.GET, queryset=queryset)
                instances = filter.qs  
                count = instances.count()

                if(instances):
                    context = self.get_context()
                    context.update({'instances': instances, 'count': count, 'filter' : filter})  
                    return render(request, self.view_all_template, context)
                else:
                    messages.add_message(request, messages.INFO, self.instance_type + " not found")
                    return redirect(request.path)    
            else:
                messages.add_message(request, messages.INFO, "Something went wrong. Invalid Search Keywordword")
                return redirect(request.path)    

        return render(request, self.form_template, context)

    def view(self, request, id):

        self.action = "View"
        
        try:
            instance = get_object_or_404(self.instance, pk=id)#, User=self.request.User)
            # instance = instance.objects.get(pk=id)   
            context = self.get_context()
            context.update({'instance': instance, 'full_information': True, 'pagesize':'A4'})  
            type = 'attachment'

            if request.method == 'POST': 
                return self.pdf_generator(request, context, type)

            return render(request, self.view_template, context)
            
        except self.instance.DoesNotExist:
            messages.add_message(request, messages.INFO, self.instance_type + " not found")
            # return redirect(self.success_url)   
            return redirect(request.META.get('HTTP_REFERER')) 
            # return redirect(instance) 
    
    def view_all(self, request):

        self.action = "View All" 
        # instance = self.instance
        try:

            # generate_random_data(self.instance, 1)
            filter = self.filter(request.GET, queryset=self.instance.objects.all())
            instances = filter.qs
            count = instances.count()  
            context = self.get_context()
            context.update({'instances': instances, 'count': count, 'filter' : filter, 'full_information': True, 'pagesize':'A4'})
            type = 'attachment'
            print(count)
            if request.method == 'POST':                 
                return self.generator(request, context, type)
                    
            return render(request, self.view_all_template, context)
        except self.instance.DoesNotExist:
            messages.add_message(request, messages.INFO, self.title + " are not found")
            # return redirect(self.success_url)   
            return redirect(request.META.get('HTTP_REFERER')) 
            # return redirect(instance) 

    
    @transaction.atomic 
    def add(self, request):
        
        self.action = "Add"   
        # instance = self.instance
        context = self.get_context()
        
        form = self.form(request.POST or None, request.FILES or None)     
        context.update({'form': form}) 

        if request.method == 'POST':

            if form.is_valid():
                try:
                    with transaction.atomic():  

                        add_form = form.save(commit=False)

    
                        try:
                            usr = User.objects.get(pk=1)
                            add_form.created_by = usr # form.created_by = request.User

                        except:                        
                            add_form.created_by =None
                        add_form.updated_by = None
                        add_form.updated_at = None
                        add_form.save()

            
                        # messages.DEBUG INFO SUCCESS WARNING ERROR
                        messages.add_message(request, messages.INFO, self.instance_type + " is Created Successfully")

                        if request.POST['submit'] == 'add':
                            return redirect(request.path) 
                        return redirect(self.success_url)
                        
                except IntegrityError:
                            messages.add_message(request, messages.INFO, self.instance_type + " cannot be created. Issue with Integrity")
                            return redirect(request.path) 
            else:
                messages.add_message(request, messages.INFO, self.instance_type + " cannot be created. Invalid Form")
                return redirect(request.path)   
        return render(request, self.form_template, context)
    
    @transaction.atomic       
    def update(self, request, id):        
        
        self.action = "Update"  
        # instance = self.instance
        context = self.get_context()

        instance = self.instance.objects.get(pk=id)
        form = self.form(request.POST or None, instance=instance)             
        context.update({'form': form})

        if request.method == 'POST':
            if form.is_valid():
                try:
                    with transaction.atomic():                 
                        form = form.save(commit=False)
                        usr = User.objects.get(pk=1)
                        form.updated_by = usr                
                        # form.updated_by = request.User
                        form.save()
                        messages.add_message(request, messages.INFO, self.instance_type + " is Updated Successfully")
                        return redirect(self.success_url)

                except IntegrityError:
                    messages.add_message(request, messages.INFO, self.instance_type + " cannot be created. Issue with Integrity")
                    return redirect(request.path)
            else:
                messages.add_message(request, messages.INFO, self.instance_type + " cannot be updated. Invalid Form")
                return redirect(request.path)   
        
        return render(request, self.form_template, context)    

    @transaction.atomic 
    def delete(self, request, id):
        
        self.action = "Delete"
        # instance = self.instance
        context = self.get_context()

        instance = self.instance.objects.get(pk=id)
        context.update({'instance': instance})

        if request.method == 'POST':     
            try:
                with transaction.atomic(): 

                    instance.delete()
                    messages.add_message(request, messages.INFO, self.instance_type + " is Deleted Successfully") 
                    return redirect(self.success_url)  

            except IntegrityError:
                messages.add_message(request, messages.INFO, self.instance_type + " cannot be deleted. Issue with Integrity")
                return redirect(request.path) 
        
        return render(request, self.assurance_template, context)
        
    def approve(self, request, id):

        self.action = "Approve"
        status = 'approved'
        return self.approve_reject(request, id, self.action, status)
        
    def reject(self, request, id):
        
        self.action = "Reject"
        status = 'rejected'
        return self.approve_reject(request, id, self.action, status)
    
    @transaction.atomic
    def approve_reject(self, request, id, action, status):        
        
        self.action = action
        # instance = self.instance
        context = self.get_context()

        instance = self.instance.objects.get(pk=id)
        context.update({'instance': instance})

        if request.method == 'POST':
            
            if instance.status == 'pending':
                try:
                    with transaction.atomic():  

                        instance.approved_by = None
                        instance.status = status
                        instance.save()
                        
                        messages.add_message(request, messages.INFO, self.instance_type + " is " + status)
                        return redirect(self.success_url)
            
                except IntegrityError:
                    messages.add_message(request, messages.INFO, self.instance_type + " cannot be " + status + ". Issue with Integrity")
                    return redirect(request.path) 
            else:
                messages.add_message(request, messages.INFO, self.instance_type + " cannot be " + status)
                return redirect(request.path)   
        
        return render(request, self.assurance_template, context)
         
    def generator(self, request, context, type):

        if request.POST['submit'] == 'preview-partial':
            type = 'inline'
            context['full_information'] = False
            
        elif request.POST['submit'] == 'preview-full':
            type = 'inline'
            
        elif request.POST['submit'] == 'download-partial-pdf':
            context['full_information'] = False

        elif request.POST['submit'] == 'download-partial-doc':
            context['full_information'] = False
            return self.generate_report_doc(request, context, type)

        elif request.POST['submit'] == 'download-full-doc':
            
            return self.generate_report_doc(request, context, type)

        # elif request.POST['submit'] == 'download-full-pdf':

        return self.generate_pdf(request, context, type)

    def generator_2(self, request, id, preference, state):
        
        instance = self.instance 
        context = self.get_context()

        instance = get_object_or_404(instance, pk=id)#, User=request.User)
        context.update({'pagesize':'A4', 'full_information': False, 'instance': instance})
        type='attachment'

        if state == 'full':
            context['full_information'] = True

        if preference == 'preview':
            type = 'inline'

        elif preference == 'download-doc':

            return self.generate_report_doc(request, context, type)


        return self.generate_pdf(request, context, type)
 
    def get_full_instance_information(self, instance, full_information):

        information = {}

        information['Code'] = str(instance.code)
        information['Start Time'] = str(instance.started_at)

        if instance.ended_at:
            information['End Time'] = str(instance.ended_at)

        information['Status'] = str(instance.status)
        information['Created By'] = str(instance.created_by)
        information['Created At'] = str(instance.created_at)

        if instance.updated_by: 
            information['Updated By'] = str(instance.updated_by) 
            information['Updated At'] = str(instance.updated_at) 

        if full_information:  

            information['Instance'] = str(instance)

            if instance.analysis_set:                
    
                extra_informations = []

                for analysis in instance.analysis_set.all(): 
                    extra_information = {}               
                    
                    if analysis.attachment_source:
                        extra_information['Source'] = str(analysis.attachment_source)                    
                        
                    

                        
                    if analysis.attachment:
                        extra_information['Attachment Data'] = analysis.attachment_data
                    
                    extra_informations.append(extra_information)

                information['Extra Informations'] = extra_informations

        return information



    def generate_pdf(self, request, context, type): 
        
        
        result = BytesIO()
        # # result = StringIO.StringIO()
        # now = str(datetime.datetime.utcnow().strftime('%H:%M:%S %d-%m-%Y')) 
        # now = ('_').join((str(datetime.datetime.now())[:-7]).split())

        now = str(datetime.datetime.now().strftime('%d-%m-%Y_%Hhr:%Mmin:%Ssec'))

        response = HttpResponse(content_type='application/pdf')
        response['Content-Disposition'] = type + '; filename=report_%s.pdf' % now
        response['Content-Encoding'] = 'UTF-8'

        html = render_to_string(self.pdf_template, context)

        # import weasyprint           
        # html = weasyprint.HTML(string=html.encode("ISO-8859-1")).write_pdf(response)
        # # import os      
        # # from django.conf import settings
        # # html = weasyprint.HTML(string=html).write_pdf(response, stylesheets=[weasyprint.CSS(os.path.join(settings.STATICFILES_DIRS[0], 'css/pdf.css'))])
        
        pisa_pdf = pisa.pisaDocument(BytesIO(html.encode("ISO-8859-1")), result)        
        if pisa_pdf.err:
            return HttpResponse('Error <pre>' + html + '</pre>' % escape(html))
        
        response.write(result.getvalue())
        result.close()

        return response

    def generate_reports_doc(self, request, context, type):
        
       headers = [
                    'No', 'ID', 'Code', 'Response Status', 'Status', 'Created At', 'Updated At','Feedback',
        ] 
        extra_headers = [
                'Attachment', 'Attachment Data',
        ]
        now = str(datetime.datetime.now().strftime('%d-%m-%Y_%Hhr:%Mmin:%Ssec'))
        today = str(datetime.datetime.now().strftime('%d-%m-%Y %H:%M:%S'))

        response = HttpResponse(content_type='application/pdf')
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = type + '; filename=report_%s.docx' % now
        response['Content-Encoding'] = 'UTF-8'

        title = context['title']
        instance_type = context['instance_type']
        pagesize = context['pagesize']
        full_information = context['full_information']

        instances = context['instances']

        document = Document()
        # document.add_page_break()
        document.add_heading(title, 0)
        document.add_paragraph('Created At: ' + today, style='ListBullet') 
        document.add_paragraph('Created By : ' + 'Khalid', style='ListBullet') 
        document.add_paragraph('Full Reports', style='ListBullet') 

        if instances:  

            informations = []
             
            for instance in instances:    

                if instance:   
                    information = self.get_full_instance_information(instance, full_information)                    
                    informations.append(information)

            # header = table.rows[0].cells
            if informations:
                
                cols = len(headers) 
                rows = len(instances)
                rows = instances.count()

                table = document.add_table(rows=rows+1, cols=cols)
                cell = table.rows[0].cells

                for j in range(cols):
                    cell[j].text = headers[j]
                
                # for j in range(rows-1):
                #     print(j+1)
                #     cell = table.rows[j+1].cells
                for i in range(rows):
                    cell = table.rows[i+1].cells
                    # cell[0].text = str(i+1)
                    information = informations[i]
                    for j in range(cols):
                        cell[j].text = information[headers[j]]
                        # print(str(headers[j+1]))
                        # print(str(information), headers)

                for i in range(rows):
                    document.add_paragraph()
                    information = informations[i]
                    # document.add_paragraph('myapp-' + str(i+1), style='ListNumber')
                    document.add_paragraph('myapp-' + str(i+1), style='IntenseQuote')
                    document = self.write_to_doc(document, information, headers, full_information)
                    
                    if i < rows-1:
                        document.add_page_break()

                    

        document.save(response)

        return response
        
    def generate_report_doc(self, request, context, type):
        
        headers = [
                    'No', 'ID', 'Code', 'Response Status', 'Status', 'Created At', 'Updated At','Feedback',
        ] 
        extra_headers = [
                'Attachment', 'Attachment Data',
        ]


       
        now = str(datetime.datetime.now().strftime('%d-%m-%Y_%Hhr:%Mmin:%Ssec'))
        today = str(datetime.datetime.now().strftime('%d-%m-%Y %H:%M:%S'))

        response = HttpResponse(content_type='application/pdf')
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = type + '; filename=report_%s.docx' % now
        response['Content-Encoding'] = 'UTF-8'

        title = context['instance_type']
        instance_type = context['instance_type']
        pagesize = context['pagesize']
        full_information = context['full_information']

        document = Document()
        document.add_heading(title + ' Report', 0)
        document.add_paragraph('Created At: ' + today, style='ListBullet') 
        document.add_paragraph('Created By : ' + 'Khalid', style='ListBullet') 

        instance = context['instance']

        if instance:
        
            information = self.get_full_instance_information(instance, full_information)  
                    
            document.add_paragraph('Full Report of '  + information['Code'], style='ListBullet') 

            document.add_paragraph('Overview -' + information['Code'], style='IntenseQuote')      
                    
            cols = 3
            rows = len(headers)

            table = document.add_table(rows=rows+1, cols=cols)
            
            cell = table.rows[0].cells
            
            cell[0].text = "No"
            cell[1].text = "Key"
            cell[2].text = "Value"
            
            for i in range(rows):
                cell = table.rows[i+1].cells
                cell[0].text = str(i+1)
                cell[1].text = headers[i]
                cell[2].text = information[headers[i]]

            document = self.write_to_doc(document, information, headers, full_information)
        else:
            document.add_paragraph('Sorry, the instance is not found', style='ListBullet') 

        document.save(response)

        return response
    
    def write_to_doc(self, document, information, headers, full_information):

        if information:              


            document.add_paragraph() 
            document.add_paragraph('Detail -' + information['Code'], style='IntenseQuote')    

            for j in range(len(headers)):
                document.add_paragraph(headers[j] + ' : ' + information[headers[j]] , style='ListBullet') 

            if full_information:   

                if 'Extra Informations' in information:  

                    extra_informations = information['Extra Informations']
                    document.add_paragraph() 
            
                    if extra_informations:
                        
                        document.add_paragraph('Analysis of myapp-' + information['Code'], style='IntenseQuote')    

                        if 'Alert Message' in information:
                            document.add_paragraph('Alert Message(Signature) :- ' + information['Alert Message'] , style='ListBullet') 
                       
                        i = 0
                        for extra_information in extra_informations:
                            i += 1 
                            # print("extra", extra_information)
                            if 'Source' in extra_information:
                                document.add_paragraph(str(i) + '. Source ' + extra_information['Source'], style='ListBullet')   

                          
                            if 'Attachment Data' in extra_information:      
                                try:
                                    document.add_picture((BytesIO(base64.decodebytes(base64.b64encode(extra_information['Attachment Data'])))), width=Inches(7))
                                    # document.add_picture((BytesIO(base64.decodebytes(bytes(information['Attachment Data'], "utf-8")))))
                                except:
                                    document.add_paragraph('No attachment found for this source' , style='ListBullet') 
                                    document.add_paragraph() 
                    else:
                        document.add_paragraph('No Analysis found-' + information['Code'], style='IntenseQuote')  
        
        # document.add_page_break()
        return document                



        # from reportlab.lib.units import mm
        # from reportlab.pdfgen.canvas import Canvas


        # class NumberedPageCanvas(Canvas):
        
        #     def __init__(self, *args, **kwargs):
        #         """Constructor"""
        #         super().__init__(*args, **kwargs)
        #         self.pages = []

        #     def showPage(self):
        #         """
        #         On a page break, add information to the list
        #         """
        #         self.pages.append(dict(self.__dict__))
        #         self._startPage()

        #     def save(self):
        #         """
        #         Add the page number to each page (page x of y)
        #         """
        #         page_count = len(self.pages)

        #         for page in self.pages:
        #             self.__dict__.update(page)
        #             self.draw_page_number(page_count)
        #             super().showPage()

        #         super().save()

        #     def draw_page_number(self, page_count):
        #         """
        #         Add the page number
        #         """
        #         page = "Page %s of %s" % (self._pageNumber, page_count)
        #         self.setFont("Helvetica", 9)
        #         self.drawRightString(179 * mm, -280 * mm, page)


        # from reportlab.pdfgen import canvas
        # from reportlab.lib.utils import ImageReader
        # cnvs = canvas.Canvas(result)

        # header_image = ImageReader('/home/khalid/myapp/Template/static/img/header.png')
        # footer_image = ImageReader('/home/khalid/myapp/Template/static/img/footer.png')
        # cnvs.drawImage(header_image, -100, 700, mask='auto')
        # cnvs.drawImage(footer_image, -100, 100, mask='auto')
        # cnvs.drawString(300, 830, "Report Created At " + now )
        # cnvs.showPage() 
        # cnvs.save() 
            # doc.build(Story, canvasmaker=PageNumCanvas)
